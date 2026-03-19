# Copyright (C) 2024-2026 Intel Corporation
# SPDX-License-Identifier: Apache-2.0

import io
import os
import re
import subprocess  # nosec B404: subprocess is used with controlled input without relaying user input to shell
import uuid
import zipfile
from typing import Optional

import docx
import pandas as pd
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from comps.text_extractor.utils.file_loaders.abstract_loader import AbstractLoader
from comps.text_extractor.utils.file_loaders.load_image import LoadImage
from comps.cores.mega.logger import get_opea_logger

logger = get_opea_logger(f"{__file__.split('comps/')[1].split('/', 1)[0]}_microservice")

# Constants
MAX_NOTE_CONTEXT_LENGTH = 100
RESERVED_NOTE_IDS = ('-1', '0')  # Word internal footnote/endnote IDs to skip

# Compiled regex patterns for XML parsing
# Matches text content in WordprocessingML: <w:t xml:space="preserve">text</w:t>
TEXT_ELEMENT_PATTERN = re.compile(r'<w:t[^>]*>([^<]*)</w:t>')
# Matches text content in DrawingML (SmartArt, diagrams): <a:t>text</a:t>
SMARTART_TEXT_PATTERN = re.compile(r'<a:t>([^<]+)</a:t>')

HEADING_STYLE_MAP = {
    'Title': '# ',
    'Heading 1': '# ',
    'Heading 2': '## ',
    'Heading 3': '### ',
    'Heading 4': '#### ',
    'Heading 5': '##### ',
    'Heading 6': '###### ',
}


class LoadDoc(AbstractLoader):
    """Load and extract text from DOC/DOCX files.
    
    Extracts:
        - Body text with heading hierarchy (markdown format)
        - Tables (pipe-delimited format)
        - Headers and footers
        - Comments (inline, after commented text)
        - Footnotes and endnotes
        - Chart data (series and values)
        - SmartArt diagram text
        - Embedded images (via OCR)
    """

    def __init__(self, file_path: str):
        """Initialize the loader.
        
        Args:
            file_path: Path to the DOC/DOCX file to process.
        """
        super().__init__(file_path)
        self._chart_data_cache: dict[str, str] = {}
        self._comments_cache: dict[str, dict] = {}

    def extract_text(self) -> str:
        """Extract all text content from the document.
        
        For .doc files, converts to .docx first using LibreOffice.
        
        Returns:
            Formatted text with sections for each content type.
        """
        original_path = self.file_path
        converted = False
        
        if self.file_type == "doc":
            self._convert_to_docx()
            converted = True

        try:
            parsed_text = self._build_document_content()
        finally:
            if converted and os.path.exists(self.file_path):
                os.remove(self.file_path)
                self.file_path = original_path
                logger.info(f"[{original_path}] Removed temporary converted file")

        return parsed_text

    def _build_document_content(self) -> str:
        """Build complete document content from all sources."""
        doc = docx.Document(self.file_path)
        logger.info(f"[{self.file_path}] Processing {len(doc.element.body)} elements")

        sections: list[str] = []

        # Extract headers/footers using official python-docx API
        headers, footers = self._extract_headers_footers(doc)

        # Extract XML-based content in single ZIP context for efficiency
        with zipfile.ZipFile(self.file_path, 'r') as zf:
            # Build relationship maps for charts and comments
            self._chart_data_cache = self._build_chart_data_cache(zf)
            self._comments_cache = self._build_comments_cache(zf)
            
            smartart = self._extract_smartart(zf)
            
            # Read document XML once for context extraction
            doc_xml = zf.read('word/document.xml').decode('utf-8')
            
            footnotes = self._extract_notes(zf, 'word/footnotes.xml', 'w:footnote', doc_xml)
            endnotes = self._extract_notes(zf, 'word/endnotes.xml', 'w:endnote', doc_xml)

            # Build output sections
            if headers:
                sections.append("[HEADER]\n" + "\n".join(headers))

            # Body extraction needs ZipFile for inline charts and comments
            body_text = self._extract_body_content(doc, zf)
            if body_text.strip():
                sections.append(body_text)

        # Clear caches
        self._chart_data_cache = {}
        self._comments_cache = {}

        if smartart:
            sections.append("[SMARTART]\n" + "\n".join(f"  {item}" for item in smartart))

        if footnotes:
            sections.append("[FOOTNOTES]\n" + "\n".join(
                self._format_note(fn) for fn in footnotes
            ))

        if endnotes:
            sections.append("[ENDNOTES]\n" + "\n".join(
                self._format_note(en) for en in endnotes
            ))

        if footers:
            sections.append("[FOOTER]\n" + "\n".join(footers))

        return "\n\n".join(sections)

    def _format_inline_comment(self, comment_id: str) -> Optional[str]:
        """Format a comment thread for inline insertion after commented text.
        
        Args:
            comment_id: The ID of the comment to format.
            
        Returns:
            Formatted comment string with replies, or None if comment not found.
        """
        if comment_id not in self._comments_cache:
            return None
        
        comment = self._comments_cache[comment_id]
        parts = [f"[{comment['author']}]: {comment['text']}"]
        
        # Add replies
        for reply in comment.get('replies', []):
            parts.append(f"[{reply['author']}]: {reply['text']}")
        
        return " [Comment: " + " | ".join(parts) + "]"

    def _format_note(self, note: dict) -> str:
        """Format a single footnote/endnote for output."""
        if note.get('context'):
            context = note['context']
            if len(context) > MAX_NOTE_CONTEXT_LENGTH:
                context = context[:MAX_NOTE_CONTEXT_LENGTH] + '...'
            return f"On \"{context}\": [{note['id']}] {note['text']}"
        return f"[{note['id']}] {note['text']}"

    def _extract_body_content(self, doc: Document, zf: zipfile.ZipFile) -> str:
        """Extract text from document body (paragraphs, tables, and inline charts)."""
        parts: list[str] = []
        for item in doc.iter_inner_content():
            if isinstance(item, docx.table.Table):
                parts.append(self._extract_table_text(item, zf))
            elif isinstance(item, docx.text.paragraph.Paragraph):
                parts.append(self._extract_paragraph_text(item, zf))
        return "".join(parts)

    def _extract_paragraph_text(self, paragraph: Paragraph, zf: zipfile.ZipFile) -> str:
        """Extract text from a paragraph, including runs, hyperlinks, and charts."""
        heading_prefix = self._get_heading_prefix(paragraph)
        para_parts: list[str] = []

        for content in paragraph.iter_inner_content():
            if isinstance(content, docx.text.run.Run):
                para_parts.append(self._extract_run_text(content, zf))
            elif isinstance(content, docx.text.hyperlink.Hyperlink):
                para_parts.append(f" {content.text} ({content.url}) ")

        text = "".join(para_parts)
        if text.strip():
            return f"{heading_prefix}{text}\n"
        return "\n" if not heading_prefix else ""

    def _extract_run_text(self, run: docx.text.run.Run, zf: zipfile.ZipFile) -> str:
        """Extract text from a run, including embedded images, charts, and inline comments."""
        parts: list[str] = []
        for content in run.iter_inner_content():
            if isinstance(content, docx.text.run.Drawing):
                drawing_text = self._process_drawing(content, zf)
                if drawing_text:
                    parts.append(f" {drawing_text} ")
            elif isinstance(content, str):
                parts.append(content)
        
        # Check for comment references in this run and append inline
        comment_refs = run._element.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference'
        )
        for ref in comment_refs:
            comment_id = ref.get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id'
            )
            if comment_id:
                inline_comment = self._format_inline_comment(comment_id)
                if inline_comment:
                    parts.append(inline_comment)
        
        return "".join(parts)

    def _process_drawing(
        self, drawing: docx.text.run.Drawing, zf: zipfile.ZipFile
    ) -> Optional[str]:
        """Process a drawing element - either chart or image."""
        # Check if this is a chart (has c:chart element)
        chart_elements = drawing._element.findall(
            './/{http://schemas.openxmlformats.org/drawingml/2006/chart}chart'
        )
        if chart_elements:
            return self._process_chart(chart_elements[0])
        
        # Otherwise treat as image
        return self._process_image(drawing)

    def _process_chart(self, chart_element) -> Optional[str]:
        """Extract chart data from cached embedded Excel spreadsheet."""
        try:
            # Get relationship ID from chart element
            r_id = chart_element.get(
                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'
            )
            if not r_id or r_id not in self._chart_data_cache:
                return None
            
            return self._chart_data_cache[r_id]
        except Exception as e:
            logger.error(f"[{self.file_path}] Error processing chart: {e}")
            return None

    def _process_image(self, drawing: docx.text.run.Drawing) -> Optional[str]:
        """Extract text from image using OCR."""
        img_path = None
        try:
            upload_path = os.getenv("UPLOAD_PATH", "/tmp/opea_upload")
            os.makedirs(upload_path, exist_ok=True)
            img_path = self._save_image(drawing, upload_path)
            if not img_path:
                return None

            img_loader = LoadImage(img_path)
            ocr_text = img_loader.extract_text().strip()
            if ocr_text:
                logger.info(f"[{self.file_path}] OCR extracted {len(ocr_text)} chars from image")
            else:
                logger.debug(f"[{self.file_path}] OCR returned no text from image")
            return ocr_text
        except Exception as e:
            logger.error(f"[{self.file_path}] Error processing image: {e}")
            return None
        finally:
            if img_path and os.path.isfile(img_path):
                os.remove(img_path)
                logger.debug(f"[{self.file_path}] Removed temp image {img_path}")

    def _extract_table_text(self, table: Table, zf: zipfile.ZipFile) -> str:
        """Extract text from a table in pipe-delimited format."""
        rows: list[str] = []
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cell_text = "".join(
                    self._extract_paragraph_text(p, zf).strip()
                    for p in cell.paragraphs
                ).strip()
                cells.append(cell_text)
            rows.append(" | ".join(cells))
        return "\n".join(rows) + "\n"

    def _get_heading_prefix(self, paragraph: Paragraph) -> str:
        """Get markdown heading prefix based on paragraph style."""
        style_name = paragraph.style.name if paragraph.style else ''
        return HEADING_STYLE_MAP.get(style_name, '')

    def _build_comments_cache(self, zf: zipfile.ZipFile) -> dict[str, dict]:
        """Build a cache mapping comment IDs to comment data with replies.
        
        Parses word/comments.xml and groups reply comments under their parent.
        Reply comments have a w:parentId attribute referencing the parent comment.
        
        Returns:
            Dict mapping comment_id -> {author, text, replies: [{author, text}, ...]}
            Only root comments (non-replies) are included as keys.
        """
        cache: dict[str, dict] = {}
        try:
            if 'word/comments.xml' not in zf.namelist():
                return cache

            comments_xml = zf.read('word/comments.xml').decode('utf-8')
            
            # Pattern matches: <w:comment w:id="X" w:author="Y" [w:parentId="Z"] ...>content</w:comment>
            comment_pattern = re.compile(
                r'<w:comment[^>]*w:id="([^"]*)"[^>]*w:author="([^"]*)"([^>]*)>(.*?)</w:comment>',
                re.DOTALL
            )
            
            # First pass: collect all comments
            all_comments: dict[str, dict] = {}
            parent_map: dict[str, str] = {}  # child_id -> parent_id
            
            for match in comment_pattern.finditer(comments_xml):
                comment_id = match.group(1)
                author = match.group(2)
                attrs = match.group(3)
                content = match.group(4)
                
                texts = TEXT_ELEMENT_PATTERN.findall(content)
                text = ' '.join(texts).strip()
                if not text:
                    continue
                
                all_comments[comment_id] = {'author': author, 'text': text, 'replies': []}
                
                # Check if this is a reply (has parentId attribute)
                parent_match = re.search(r'w:parentId="([^"]*)"', attrs)
                if parent_match:
                    parent_map[comment_id] = parent_match.group(1)
            
            # Second pass: group replies under parents
            for comment_id, data in all_comments.items():
                if comment_id in parent_map:
                    # its reply -> add to parents replies list
                    parent_id = parent_map[comment_id]
                    if parent_id in all_comments:
                        all_comments[parent_id]['replies'].append({
                            'author': data['author'],
                            'text': data['text']
                        })
                else:
                    # root comment - add to cache
                    cache[comment_id] = data
            
            if cache:
                total_replies = sum(len(c['replies']) for c in cache.values())
                logger.info(
                    f"[{self.file_path}] Cached {len(cache)} comments with {total_replies} replies"
                )
        except Exception as e:
            logger.error(f"[{self.file_path}] Error building comments cache: {e}")

        return cache

    def _extract_notes(
        self, zf: zipfile.ZipFile, xml_path: str, tag_name: str, doc_xml: str
    ) -> list[dict]:
        """Extract footnotes or endnotes from XML with their context.
        
        Args:
            zf: Open ZipFile containing the DOCX
            xml_path: Path to XML file (e.g., 'word/footnotes.xml')
            tag_name: XML tag name (e.g., 'w:footnote')
            doc_xml: Document XML content for context extraction
        """
        notes: list[dict] = []
        try:
            if xml_path not in zf.namelist():
                return notes

            xml = zf.read(xml_path).decode('utf-8')
            # Match note elements: <w:footnote w:id="1">...</w:footnote>
            pattern = re.compile(
                rf'<{tag_name}[^>]*w:id="(-?\d+)"[^>]*>(.*?)</{tag_name}>',
                re.DOTALL
            )
            
            # Determine reference tag name (footnoteReference or endnoteReference)
            ref_tag = 'footnoteReference' if 'footnote' in tag_name else 'endnoteReference'

            for match in pattern.finditer(xml):
                note_id = match.group(1)
                if note_id in RESERVED_NOTE_IDS:
                    continue

                content = match.group(2)
                texts = TEXT_ELEMENT_PATTERN.findall(content)
                text = ' '.join(texts).strip()
                if text:
                    # Find context - the paragraph containing this note reference
                    context = self._find_note_context(doc_xml, note_id, ref_tag)
                    notes.append({'id': note_id, 'text': text, 'context': context})

            if notes:
                logger.info(f"[{self.file_path}] Extracted {len(notes)} notes from {xml_path}")
        except Exception as e:
            logger.error(f"[{self.file_path}] Error extracting notes from {xml_path}: {e}")

        return notes

    def _find_note_context(
        self, doc_xml: str, note_id: str, ref_tag: str
    ) -> Optional[str]:
        """Find the paragraph text containing a footnote/endnote reference.
        
        Args:
            doc_xml: Document XML content
            note_id: ID of the note to find
            ref_tag: Reference tag name ('footnoteReference' or 'endnoteReference')
        """
        try:
            # Find paragraphs containing the note reference
            para_pattern = r'<w:p[^>]*>(.*?)</w:p>'
            ref_pattern = rf'<w:{ref_tag}[^/]*w:id="{note_id}"'
            
            for para_match in re.finditer(para_pattern, doc_xml, re.DOTALL):
                para_content = para_match.group(1)
                if re.search(ref_pattern, para_content):
                    # Extract text from this paragraph
                    texts = TEXT_ELEMENT_PATTERN.findall(para_content)
                    context = ' '.join(texts).strip()
                    return context if context else None
        except Exception as e:
            logger.debug(f"[{self.file_path}] Could not find context for note {note_id}: {e}")
        return None

    def _extract_headers_footers(
        self, doc: Document
    ) -> tuple[list[str], list[str]]:
        """Extract text from document headers and footers using python-docx API.
        
        Each section can have its own header/footer, or inherit from previous.
        We collect unique header/footer texts across all sections.
        """
        headers: list[str] = []
        footers: list[str] = []

        try:
            for section in doc.sections:
                # Header - only if this section has its own definition
                if not section.header.is_linked_to_previous:
                    text = self._extract_header_footer_text(section.header)
                    if text and text not in headers:
                        headers.append(text)
                
                # Footer - only if this section has its own definition
                if not section.footer.is_linked_to_previous:
                    text = self._extract_header_footer_text(section.footer)
                    if text and text not in footers:
                        footers.append(text)

            if headers or footers:
                logger.info(
                    f"[{self.file_path}] Extracted {len(headers)} headers, {len(footers)} footers"
                )
        except Exception as e:
            logger.error(f"[{self.file_path}] Error extracting headers/footers: {e}")

        return headers, footers

    def _extract_header_footer_text(self, header_footer) -> str:
        """Extract text from a header or footer object."""
        texts = [p.text.strip() for p in header_footer.paragraphs if p.text.strip()]
        return '\n'.join(texts)

    def _build_chart_data_cache(
        self, zf: zipfile.ZipFile
    ) -> dict[str, str]:
        """Build a cache mapping chart relationship IDs to formatted chart data.
        
        DOCX chart data flow:
        1. document.xml has <c:chart r:id="rIdX"/>
        2. document.xml.rels maps rIdX -> charts/chart1.xml
        3. chart1.xml.rels maps to ../embeddings/Microsoft_Excel_Worksheet.xlsx
        4. Excel file contains the actual data
        
        Returns:
            Dict mapping relationship ID (e.g., 'rId15') to formatted chart text.
        """
        cache: dict[str, str] = {}
        try:
            # Parse document.xml.rels to get chart relationship IDs
            if 'word/_rels/document.xml.rels' not in zf.namelist():
                return cache
            
            doc_rels = zf.read('word/_rels/document.xml.rels').decode('utf-8')
            
            # Find all chart relationships: rId -> charts/chartN.xml
            chart_rel_pattern = re.compile(
                r'Relationship[^>]*Id="([^"]+)"[^>]*Target="(charts/chart\d+\.xml)"'
            )
            
            for match in chart_rel_pattern.finditer(doc_rels):
                r_id = match.group(1)
                chart_path = f"word/{match.group(2)}"
                
                # Get Excel path from charts relationships
                chart_rels_path = chart_path.replace('.xml', '.xml.rels')
                chart_rels_path = chart_rels_path.replace('word/charts/', 'word/charts/_rels/')
                
                if chart_rels_path not in zf.namelist():
                    continue
                
                chart_rels = zf.read(chart_rels_path).decode('utf-8')
                excel_match = re.search(r'Target="([^"]*\.xlsx)"', chart_rels)
                
                if not excel_match:
                    continue
                
                # Resolve relative path (../embeddings/file.xlsx)
                excel_rel_path = excel_match.group(1)
                if excel_rel_path.startswith('..'):
                    excel_path = f"word/{excel_rel_path.lstrip('../')}"
                else:
                    excel_path = f"word/charts/{excel_rel_path}"
                
                if excel_path not in zf.namelist():
                    continue
                
                # Parse Excel and cache result
                xlsx_data = zf.read(excel_path)
                chart_text = self._parse_embedded_excel(xlsx_data)
                if chart_text:
                    cache[r_id] = chart_text
                    
            if cache:
                logger.info(f"[{self.file_path}] Cached {len(cache)} chart(s) data")
                
        except Exception as e:
            logger.error(f"[{self.file_path}] Error building chart cache: {e}")
        
        return cache

    def _parse_embedded_excel(self, xlsx_data: bytes) -> Optional[str]:
        """Parse embedded Excel file and format as chart data.
        
        Args:
            xlsx_data: Raw Excel file bytes.
            
        Returns:
            Formatted chart data as string, or None if parsing fails.
        """
        try:
            df = pd.read_excel(io.BytesIO(xlsx_data))
            
            if df.empty:
                return None
            
            # Format as readable table
            lines = ["[CHART]"]
            lines.append(df.to_string(index=False))
            
            return "\n".join(lines)
        except Exception as e:
            logger.debug(f"[{self.file_path}] Could not parse embedded Excel: {e}")
            return None

    def _extract_smartart(self, zf: zipfile.ZipFile) -> list[str]:
        """Extract text from SmartArt diagrams."""
        texts: list[str] = []
        try:
            for name in zf.namelist():
                if name.startswith('word/diagrams/data') and name.endswith('.xml'):
                    xml = zf.read(name).decode('utf-8')
                    for match in SMARTART_TEXT_PATTERN.finditer(xml):
                        entry = match.group(1).strip()
                        if entry and entry != '[Text]':
                            texts.append(entry)

            if texts:
                logger.info(f"[{self.file_path}] Extracted {len(texts)} SmartArt elements")
        except Exception as e:
            logger.error(f"[{self.file_path}] Error extracting SmartArt: {e}")

        return texts

    def _save_image(
        self, drawing: docx.text.run.Drawing, save_path: str
    ) -> Optional[str]:
        """Save embedded image from drawing element to disk.
        
        Args:
            drawing: Drawing element containing the image.
            save_path: Directory path to save the image.
            
        Returns:
            Full path to saved image file, or None if no image found.
        """
        blip = drawing._element.xpath('.//a:blip')
        if not blip:
            return None

        r_embed = blip[0].get(
            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
        )
        image_part = drawing._parent.part.related_parts[r_embed]
        image_ext = image_part.content_type.split('/')[-1]
        image_filename = os.path.join(save_path, f"{uuid.uuid4()}.{image_ext}")

        with open(image_filename, 'wb') as f:
            f.write(image_part.blob)

        logger.debug(f"[{self.file_path}] Saved image: {image_filename}")
        return image_filename

    def _convert_to_docx(self) -> None:
        """Convert .doc file to .docx using LibreOffice.
        
        Updates self.file_path to point to the converted file.
        
        Raises:
            ValueError: If conversion fails.
        """
        doc_path = self.file_path
        output_dir = os.path.dirname(doc_path)
        docx_path = doc_path + "x"

        try:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--invisible",
                    "--convert-to", "docx",
                    "--outdir", output_dir,
                    doc_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0 or not os.path.exists(docx_path):
                error_msg = result.stderr or result.stdout or "Unknown error"
                logger.error(
                    f"[{self.file_path}] LibreOffice conversion failed: {error_msg}"
                )
                raise ValueError(
                    f"Failed to convert {doc_path} to docx format. Error: {error_msg}"
                )
                
        except subprocess.TimeoutExpired:
            logger.error(f"[{self.file_path}] LibreOffice conversion timed out")
            raise ValueError(f"Conversion of {doc_path} timed out after 60 seconds")

        self.file_path = docx_path
        self.file_type = "docx"
        logger.info(f"[{doc_path}] Converted to docx")
